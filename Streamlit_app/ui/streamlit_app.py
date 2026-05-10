import os
import sys
import re
import base64
from io import BytesIO
from datetime import datetime
from typing import cast
from urllib.parse import quote_plus

import streamlit as st
from docx import Document
from pypdf import PdfReader


# -------------------------------------------------
# Path Setup
# -------------------------------------------------

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
ROOT_DIR = os.path.abspath(os.path.join(BASE_DIR, ".."))

sys.path.append(ROOT_DIR)

from langgraph_app import careerskills_app, CareerSkillsState
from utils.email_helper import send_acknowledgement_email

# -------------------------------------------------
# Asset Paths
# -------------------------------------------------

LOGO_PATH = os.path.join(BASE_DIR, "static", "images", "careerskillsai_logo.png")
BANNER_PATH = os.path.join(BASE_DIR, "static", "images", "careerskillsai_banner.png")
ICON_DIR = os.path.join(BASE_DIR, "static", "icons")
CSS_PATH = os.path.join(BASE_DIR, "static", "css", "streamlit_style.css")

# -------------------------------------------------
# Helper Functions
# -------------------------------------------------

def get_base64(file_path):
    with open(file_path, "rb") as f:
        return base64.b64encode(f.read()).decode()


def read_uploaded_file(uploaded_file):
    if uploaded_file is None:
        return ""

    file_name = uploaded_file.name.lower()

    if file_name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")

    if file_name.endswith(".docx"):
        doc = Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])

    if file_name.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    return ""


def extract_final_resume(ai_response):
    marker = "## Final Tailored Resume Draft"

    if marker in ai_response:
        return ai_response.split(marker, 1)[1].strip()

    return ai_response


def create_docx_download(text):
    doc = Document()
    doc.add_heading("Final Tailored Resume Draft", level=1)

    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def add_course_search_links(ai_response):
    pattern = r'Search Keyword To Use:\s*"?([^"\n]+)"?'

    def replace_match(match):
        keyword = match.group(1).strip()

        url = (
            "https://courses.myskillsfuture.gov.sg/search?q="
            + quote_plus(keyword)
            + "&page=1&hasUpcomingCourseRun=false"
        )

        return (
            f"Search Keyword To Use: {keyword}\n\n"
            f"[Open MySkillsFuture Search]({url})"
        )

    return re.sub(pattern, replace_match, ai_response)


def run_agent(
    query,
    resume_text="",
    job_description="",
    job_title="",
    target_role="",
    career_goal="",
    resume_mode="advice"
):
    state = {
        "query": query,
        "resume_text": resume_text,
        "job_description": job_description,
        "job_title": job_title,
        "target_role": target_role,
        "career_goal": career_goal,
        "resume_mode": resume_mode,
        "next_node": "",
        "response": ""
    }

    return careerskills_app.invoke(cast(CareerSkillsState, state))

def load_css(file_path):
    with open(file_path) as f:
        st.markdown(
            f"<style>{f.read()}</style>",
            unsafe_allow_html=True
        )

# -------------------------------------------------
# Load css
# -------------------------------------------------

load_css(CSS_PATH)

# -------------------------------------------------
# Page Config
# -------------------------------------------------

st.set_page_config(
    page_title="Career&Skills AI",
    page_icon="🚀",
    layout="wide"
)

# -------------------------------------------------
# Sidebar
# -------------------------------------------------

with st.sidebar:
    st.image(LOGO_PATH, use_container_width=True)

    st.markdown("### Career Navigation Services")

    service = st.radio(
        "Select a service",
        [
            "🏠 Home",
            "📄 Resume Rewrite",
            "🎓 Course Recommendation",
            "💼 Career / Job Advice",
            "🎤 Interview Preparation",
            "🛡️ Policy & Privacy",
            "👤 Human Escalation"
        ]
    )


# -------------------------------------------------
# Main Banner
# -------------------------------------------------

# st.image(BANNER_PATH, use_container_width=True)

col1, col2, col3 = st.columns([1, 8, 1])

with col2:
    st.markdown('<div class="banner-img">', unsafe_allow_html=True)

    st.image(
        BANNER_PATH,
        use_container_width=True
    )

    st.markdown('</div>', unsafe_allow_html=True)

# -------------------------------------------------
# Home Page
# -------------------------------------------------

if service == "🏠 Home":
    st.markdown("## 👋 Welcome to Career&Skills AI")

    st.markdown("""
    Career&Skills AI is an AI-powered career and upskilling assistant that helps users rewrite resumes,
    discover relevant courses, explore job pathways, prepare for interviews, and request human support.
    """)

    st.markdown("### 🌟 Services")

    def service_card(icon_file, title, description):
        icon_path = os.path.join(ICON_DIR, icon_file)
        icon_base64 = get_base64(icon_path)

        st.markdown(f"""
        <div class="service-card">
            <img src="data:image/png;base64,{icon_base64}" width="90">
            <div class="service-title">{title}</div>
            <p>{description}</p>
        </div>
        """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)

    with col1:
        service_card(
            "icon_resume.png",
            "Resume Rewrite",
            "Rewrite and improve your resume for a specific job application."
        )

    with col2:
        service_card(
            "icon_course.png",
            "Course Recommendation",
            "Find SkillsFuture-related courses based on your goals and skill gaps."
        )

    with col3:
        service_card(
            "icon_career.png",
            "Career / Job Advice",
            "Explore career pathways, suitable roles, and job-related skill requirements."
        )

    col4, col5, col6 = st.columns(3)

    with col4:
        service_card(
            "icon_interview.png",
            "Interview Preparation",
            "Practise role-specific interview questions and preparation guidance."
        )

    with col5:
        service_card(
            "icon_policy.png",
            "Policy & Privacy",
            "Understand safe system usage and what personal data not to upload."
        )

    with col6:
        service_card(
            "icon_human.png",
            "Human Escalation",
            "Submit a human review request with ticket number and email acknowledgement."
        )


# -------------------------------------------------
# Resume Rewrite
# -------------------------------------------------

elif service == "📄 Resume Rewrite":
    st.markdown("## 📄 Resume Rewrite Agent")

    st.info("""
    Rewrite and optimize your resume for a specific job application.

    This service improves professionalism, ATS-readability, formatting, and keyword alignment
    without adding fake skills or experience.
    """)

    st.warning("""
    Before pasting or uploading your resume, please remove sensitive personal information such as:
    NRIC, passport number, mobile number, home address, date of birth, bank details, passwords,
    and private identification numbers.
    """)

    query = st.text_area(
        "Your Question / Instruction",
        value="Please rewrite my resume for this job application. Make it professional, ATS-friendly, and aligned to the job description without adding skills I do not have."
    )

    uploaded_resume = st.file_uploader(
        "Upload Resume (.txt, .docx, .pdf)",
        type=["txt", "docx", "pdf"]
    )

    resume_from_file = read_uploaded_file(uploaded_resume)

    resume_text = st.text_area(
        "Paste Resume Content",
        value=resume_from_file,
        height=260
    )

    job_title = st.text_input(
        "Job Title Apply",
        placeholder="Example: Data Analyst / AI Engineer / Business Analyst"
    )

    job_description = st.text_area(
        "Paste Job Description",
        height=260
    )

    if st.button("Submit"):
        with st.spinner("Career&Skills AI is rewriting your resume..."):
            result = run_agent(
                query="resume rewrite: " + query,
                resume_text=resume_text,
                job_description=job_description,
                job_title=job_title
            )

            st.success(f"Selected Agent: {result['next_node']}")
            st.markdown(result["response"])

            final_resume = extract_final_resume(result["response"])
            docx_file = create_docx_download(final_resume)

            st.download_button(
                label="Download Final Tailored Resume Draft",
                data=docx_file,
                file_name="final_tailored_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


# -------------------------------------------------
# Course Recommendation
# -------------------------------------------------

elif service == "🎓 Course Recommendation":
    st.markdown("## 🎓 Course Recommendation Agent")

    st.info("""
    Discover Singapore-relevant upskilling and career transition courses based on your current skills,
    target role, and career goals.

    Career&Skills AI primarily recommends SkillsFuture / MySkillsFuture-related courses and professional
    training programs to support career growth, job transition, and employability.
    """)

    query = st.text_area(
        "Your Question / Instruction",
        value="Recommend suitable courses for my career goal."
    )

    resume_text = st.text_area(
        "Paste Resume / Skill Summary",
        height=220
    )

    target_role = st.text_input(
        "Target Role",
        placeholder="Example: AI Engineer, Data Analyst, Finance Analyst"
    )

    career_goal = st.text_area(
        "Career Goal",
        placeholder="Example: I want to upskill from Data Analyst to AI Engineer."
    )

    if st.button("Submit"):
        with st.spinner("Career&Skills AI is recommending courses..."):
            result = run_agent(
                query="course recommendation: " + query,
                resume_text=resume_text,
                target_role=target_role,
                career_goal=career_goal
            )

            st.success(f"Selected Agent: {result['next_node']}")
            response_with_links = add_course_search_links(result["response"])
            st.markdown(response_with_links)


# -------------------------------------------------
# Career / Job Advice
# -------------------------------------------------

elif service == "💼 Career / Job Advice":
    st.markdown("## 💼 Career / Job Advice Agent")

    st.info("""
    Learn about job roles, required skills, career pathways, and suitable opportunities.

    Useful for users exploring career transitions, new industries, or understanding what skills employers are looking for.
    """)

    query = st.text_area(
        "Your Question / Instruction",
        placeholder="Example: What jobs are suitable for someone with Python, SQL and Excel?"
    )

    resume_text = st.text_area(
        "Paste Resume / Skill Summary Optional",
        height=220
    )

    if st.button("Submit"):
        with st.spinner("Career&Skills AI is checking suitable career options..."):
            result = run_agent(
                query="career recommendation and job matching advice: " + query,
                resume_text=resume_text
            )

            st.success(f"Selected Agent: {result['next_node']}")
            st.markdown(result["response"])


# -------------------------------------------------
# Interview Preparation
# -------------------------------------------------

elif service == "🎤 Interview Preparation":
    st.markdown("## 🎤 Interview Preparation Agent")

    st.info("""
    Prepare for interviews with AI-generated mock interview questions, suggested answer guidance,
    communication tips, and role-specific preparation advice.
    """)

    query = st.text_area(
        "Your Question / Instruction\n "
        "(Optional: may want to add in the Job Description or specific areas you want to prepare for, e.g. technical questions, behavioural questions, or both) below your question",
        value="Prepare me for interview."
    )

    target_role = st.text_input(
        "Target Role",
        placeholder="Example: Data Analyst"
    )

    resume_text = st.text_area(
        "Paste Background / Resume Summary Optional",
        height=220
    )

    if st.button("Submit"):
        with st.spinner("Career&Skills AI is preparing interview questions..."):
            result = run_agent(
                query="interview preparation: " + query,
                resume_text=resume_text,
                target_role=target_role
            )

            st.success(f"Selected Agent: {result['next_node']}")
            st.markdown(result["response"])


# -------------------------------------------------
# Policy & Privacy
# -------------------------------------------------

elif service == "🛡️ Policy & Privacy":
    st.markdown("## 🛡️ Policy & Privacy")

    st.info("""
    Understand the system's privacy policy, responsible AI usage guidelines, and important data protection reminders.
    Users are strongly advised not to upload confidential or sensitive personal information into the system.
    """)

    st.error("""
    Do NOT upload sensitive personal information:
    - NRIC / passport number
    - Bank account details
    - Passwords
    - Home address
    - Date of birth
    - Mobile number
    - Private identification numbers
    """)

    st.markdown("""
    ### How to use Career&Skills AI safely

    1. Remove personal details before uploading your resume.
    2. Keep only career-related information.
    3. Use sample or anonymised data for testing.
    4. Verify AI recommendations before making career decisions.
    5. Escalate complex career cases to a human advisor.

    ### System Policy

    Career&Skills AI is an educational AI career assistant.  
    It provides resume, career, course, and interview guidance.  
    It does not replace professional career counselling or legal advice.
    """)


# -------------------------------------------------
# Human Escalation
# -------------------------------------------------

elif service == "👤 Human Escalation":
    st.markdown("## 👤 Human Escalation")

    st.info("""
    Submit a request for human review when your situation requires personal career guidance,
    resume review, career transition advice, or support beyond AI-generated recommendations.

    This is a prototype workflow that generates a ticket number and sends an acknowledgement email.
    """)

    title = st.selectbox(
        "Title / Salutation",
        [
            "Mr.",
            "Ms.",
            "Mrs.",
            "Mdm.",
            "Dr.",
            "Prefer not to say"
        ]
    )

    col1, col2 = st.columns(2)

    with col1:
        first_name = st.text_input("First Name")

    with col2:
        last_name = st.text_input("Last Name")

    years_experience = st.number_input(
        "Years of Working Experience",
        min_value=0,
        max_value=60,
        step=1
    )

    email = st.text_input("Email Address")

    uploaded_resume = st.file_uploader(
        "Upload Resume (Optional)",
        type=["txt", "docx", "pdf"]
    )

    request_description = st.text_area(
        "Describe What You Need",
        height=180,
        placeholder="Example: I need a human advisor to review my resume for a Data Analyst job."
    )

    if st.button("Submit"):
        ticket_id = "CSAI-" + datetime.now().strftime("%Y%m%d%H%M%S")

        if not first_name or not last_name or not email or not request_description:
            st.error("Please fill in First Name, Last Name, Email, and Description.")
        else:
            success, message = send_acknowledgement_email(
                to_email=email,
                ticket_id=ticket_id,
                title=title,
                first_name=first_name,
                request_description=request_description
            )

            st.success(f"Human escalation request submitted. Ticket Number: {ticket_id}")

            if success:
                st.info("Acknowledgement email sent successfully.")
            else:
                st.warning(f"Ticket created, but email was not sent: {message}")